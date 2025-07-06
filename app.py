import streamlit as st
import pandas as pd
from docx import Document
from datetime import date
import tempfile
import base64
import fitz  # PyMuPDF

def analyze_market(df):
    # Required columns
    required_columns = ["close price", "above grade finished area", "concessions", "address", "bedrooms total", "bathrooms total integer"]
    for col in required_columns:
        if col not in df.columns:
            st.error(f"Your data must include a '{col}' column.")
            st.stop()

    comps = df.copy()

    # Fill missing concessions
    comps["concessions"] = comps["concessions"].fillna(0)

    # Remove comps without valid square footage
    comps = comps[comps["above grade finished area"] > 0]
    if comps.empty:
        st.error("No comps have valid square footage.")
        st.stop()

    # Compute net price
    comps["NetPrice"] = comps["close price"] - comps["concessions"]

    # Compute PPSF
    comps["PricePerSF"] = comps["NetPrice"] / comps["above grade finished area"]

    avg_ppsf = comps["PricePerSF"].mean()

    return comps, avg_ppsf

def extract_pdf_text(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    all_text = []
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text()
        all_text.append(f"--- Subject Property Details - Page {page_num} ---\n{text}\n")
    return "\n".join(all_text)

def generate_report(subject_info, comps, est_ppsf, notes, zillow_val, redfin_val, pdf_text):
    doc = Document()
    doc.add_heading('Market Valuation Report', 0)
    doc.add_paragraph(f"Date: {date.today().strftime('%B %d, %Y')}")

    doc.add_heading('Subject Property', level=1)
    doc.add_paragraph(f"Address: {subject_info['address']}")
    doc.add_paragraph(
        f"SqFt: {subject_info['sqft']}, Beds: {subject_info['beds']}, Baths: {subject_info['baths']}"
    )
    doc.add_paragraph(f"Close Price: ${subject_info['price']:,.0f}")

    est_subject_price = est_ppsf * subject_info['sqft']
    doc.add_heading('Estimated Market Value Based on Comps', level=1)
    doc.add_paragraph(f"Average Net Price per SqFt: ${est_ppsf:,.2f}")
    doc.add_paragraph(f"Estimated Value: ${est_subject_price:,.0f}")

    doc.add_heading('Public Online Estimates', level=1)
    doc.add_paragraph(f"Zillow Zestimate: ${zillow_val}")
    doc.add_paragraph(f"Redfin Estimate: ${redfin_val}")

    doc.add_heading('Notes and Special Features', level=1)
    doc.add_paragraph(notes)

    doc.add_heading('Comparable Properties (Net Prices)', level=1)
    for _, row in comps.iterrows():
        doc.add_paragraph(
            f"{row['address']}: Net ${row['NetPrice']:,.0f} | {row['above grade finished area']} SqFt | {row['bedrooms total']} Bd / {row['bathrooms total integer']} Ba | PPSF ${row['PricePerSF']:,.2f}",
            style='List Bullet'
        )

    if pdf_text:
        doc.add_heading('Subject Property Details (from PDF)', level=1)
        doc.add_paragraph(pdf_text)

    doc.add_paragraph("\n---\n")
    doc.add_paragraph(
        "This is an estimate based on MLS market information and publicly available data. "
        "It is intended for marketing and informational purposes only and does not constitute an appraisal or guarantee of market value."
    )

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

st.title("📊 Cloud Market Valuation Tool")
st.write("Upload MLS data, enter subject details, and generate a branded market report.")

uploaded_csv = st.file_uploader("Upload MLS CSV/XLSX File", type=["csv", "xlsx"])
uploaded_pdf = st.file_uploader("Upload Subject Property PDF (optional)", type=["pdf"])

if uploaded_csv:
    # Load CSV or XLSX
    if uploaded_csv.name.endswith(".csv"):
        df = pd.read_csv(uploaded_csv)
    else:
        df = pd.read_excel(uploaded_csv)

    # Normalize columns
    df.columns = [col.strip().lower() for col in df.columns]

    # Build Address column from parts
    df["address"] = (
        df["street number"].astype(str).str.strip()
        + " "
        + df["street dir prefix"].fillna("").str.strip() + " "
        + df["street name"].str.strip() + " "
        + df["street dir suffix"].fillna("").str.strip()
    ).str.replace("  ", " ").str.strip()

    st.write("Columns detected:", df.columns.tolist())

    comps, avg_ppsf = analyze_market(df)

    st.write("Preview of Comparable Properties with Net Prices:")
    st.dataframe(comps[["address", "NetPrice", "PricePerSF", "above grade finished area", "bedrooms total", "bathrooms total integer"]])

    st.subheader("Enter Subject Property Details")

    subject_address = st.text_input("Subject Property Address")
    subject_sqft = st.number_input("Above Grade Finished Area (SqFt)", min_value=0)
    subject_beds = st.number_input("Bedrooms", min_value=0, step=1)
    subject_baths = st.number_input("Bathrooms", min_value=0, step=1)
    subject_price = st.number_input("Close Price", min_value=0)

    notes = st.text_area("Notes and Special Features")
    zillow_val = st.text_input("Zillow Zestimate")
    redfin_val = st.text_input("Redfin Estimate")

    if st.button("Generate Report"):
        subject_info = {
            "address": subject_address,
            "sqft": subject_sqft,
            "beds": subject_beds,
            "baths": subject_baths,
            "price": subject_price
        }

        pdf_text = ""
        if uploaded_pdf:
            pdf_text = extract_pdf_text(uploaded_pdf)

        report_path = generate_report(subject_info, comps, avg_ppsf, notes, zillow_val, redfin_val, pdf_text)

        with open(report_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
            href = f'<a href="data:application/octet-stream;base64,{b64}" download="Market_Valuation_Report.docx">📄 Download Report</a>'
            st.markdown(href, unsafe_allow_html=True)


